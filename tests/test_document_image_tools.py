from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import pytest
from docx import Document
from PIL import Image
from pypdf import PdfWriter

from mai.agent import WorkContext
from mai.document_tools import DocumentReadTool, ImageAnalyzeTool, build_document_image_tools
from mai.file_tools import FileToolAccess, FileToolAuthorizationError


def context(user_id: str = "owner") -> WorkContext:
    return WorkContext(user_id=user_id, turn_id="turn", user_text="test")


@dataclass
class FakeAnalyzer:
    model: str = "gemma4:12b"
    calls: list[dict] = field(default_factory=list)

    def analyze(self, *, path: Path, prompt: str) -> str:
        self.calls.append({"path": path, "prompt": prompt})
        return "image result"


def access(tmp_path: Path) -> FileToolAccess:
    return FileToolAccess(owner_id="owner", default_root=tmp_path)


def test_builds_exact_document_and_image_tool_names(tmp_path) -> None:
    tools = build_document_image_tools(owner_id="owner", analyzer=FakeAnalyzer(), default_root=tmp_path)
    assert [tool.name for tool in tools] == ["document_read", "image_analyze"]


def test_document_read_schema_only_accepts_pdf_or_docx_paths(tmp_path) -> None:
    schema = DocumentReadTool(access(tmp_path)).schema()
    path_schema = schema["properties"]["arguments"]["properties"]["path"]
    assert path_schema["pattern"] == r".*\.(?:[pP][dD][fF]|[dD][oO][cC][xX])$"


def test_document_read_docx_uses_paragraph_pagination(tmp_path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("one")
    document.add_paragraph("two")
    document.add_paragraph("three")
    document.save(path)

    result = DocumentReadTool(access(tmp_path)).execute(
        arguments={"path": str(path), "start": 2, "limit": 1},
        context=context(),
    )

    assert result["document_type"] == "docx"
    assert result["unit"] == "paragraph"
    assert result["items"] == [{"paragraph": 2, "text": "two"}]
    assert result["total"] == 3
    assert result["has_more"] is True
    assert result["next_start"] == 3


def test_document_read_pdf_uses_page_pagination(tmp_path) -> None:
    path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.add_blank_page(width=72, height=72)
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as handle:
        writer.write(handle)

    result = DocumentReadTool(access(tmp_path)).execute(
        arguments={"path": str(path), "start": 2, "limit": 1},
        context=context(),
    )

    assert result["document_type"] == "pdf"
    assert result["unit"] == "page"
    assert [item["page"] for item in result["items"]] == [2]
    assert result["total"] == 3
    assert result["next_start"] == 3


def test_document_read_rejects_unsupported_type_before_path_existence(tmp_path) -> None:
    path = tmp_path / "missing.txt"
    assert not path.exists()
    with pytest.raises(ValueError, match="unsupported document type"):
        DocumentReadTool(access(tmp_path)).execute(arguments={"path": str(path)}, context=context())


def test_document_read_missing_supported_document_still_fails_as_missing(tmp_path) -> None:
    path = tmp_path / "missing.pdf"
    with pytest.raises(FileNotFoundError):
        DocumentReadTool(access(tmp_path)).execute(arguments={"path": str(path)}, context=context())


def test_document_and_image_tools_are_owner_only(tmp_path) -> None:
    path = tmp_path / "sample.docx"
    Document().save(path)
    with pytest.raises(FileToolAuthorizationError):
        DocumentReadTool(access(tmp_path)).execute(arguments={"path": str(path)}, context=context("member"))

    image_path = tmp_path / "sample.png"
    Image.new("RGB", (2, 2)).save(image_path)
    with pytest.raises(FileToolAuthorizationError):
        ImageAnalyzeTool(access(tmp_path), FakeAnalyzer()).execute(
            arguments={"path": str(image_path), "prompt": "describe"},
            context=context("member"),
        )


def test_image_analyze_validates_image_and_uses_independent_analyzer(tmp_path) -> None:
    path = tmp_path / "sample.png"
    Image.new("RGB", (2, 2)).save(path)
    analyzer = FakeAnalyzer()

    result = ImageAnalyzeTool(access(tmp_path), analyzer).execute(
        arguments={"path": str(path), "prompt": "what is shown?"},
        context=context(),
    )

    assert result == {"path": str(path.resolve()), "model": "gemma4:12b", "analysis": "image result"}
    assert analyzer.calls == [{"path": path.resolve(), "prompt": "what is shown?"}]


def test_image_analyze_invalid_image_failure_is_visible(tmp_path) -> None:
    path = tmp_path / "fake.png"
    path.write_bytes(b"not an image")
    with pytest.raises(Exception):
        ImageAnalyzeTool(access(tmp_path), FakeAnalyzer()).execute(
            arguments={"path": str(path), "prompt": "describe"},
            context=context(),
        )
