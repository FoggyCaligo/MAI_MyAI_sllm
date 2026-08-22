from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from PIL import Image
from pypdf import PdfReader

from .document_tools import ImageAnalyzer


_TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".py", ".js", ".ts", ".tsx", ".jsx", ".json", ".yaml", ".yml",
    ".toml", ".ini", ".cfg", ".conf", ".csv", ".tsv", ".html", ".htm", ".css", ".sql", ".xml",
    ".sh", ".ps1", ".bat", ".cmd", ".java", ".kt", ".kts", ".c", ".h", ".cpp", ".hpp", ".cs",
    ".go", ".rs", ".rb", ".php", ".swift", ".vue", ".svelte",
}
_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tif", ".tiff"}


@dataclass(slots=True)
class AttachmentEvidenceBuilder:
    analyzer: ImageAnalyzer
    text_encoding: str = "utf-8"
    per_item_chars: int = 8000
    total_chars: int = 24000

    def build(self, paths: Iterable[str | Path]) -> list[dict[str, Any]]:
        remaining = max(0, int(self.total_chars))
        evidence: list[dict[str, Any]] = []
        for index, raw_path in enumerate(paths, start=1):
            path = Path(raw_path).expanduser().resolve()
            source_id = f"attachment:{index}"
            if not path.exists():
                raise FileNotFoundError(path)
            if not path.is_file():
                raise IsADirectoryError(path)
            if remaining <= 0:
                evidence.append(
                    {
                        "evidence_id": source_id,
                        "path": str(path),
                        "status": "not_loaded_context_budget",
                    }
                )
                continue
            item = self._read(path=path, evidence_id=source_id, char_limit=min(self.per_item_chars, remaining))
            evidence.append(item)
            remaining -= len(str(item.get("content") or ""))
        return evidence

    def _read(self, *, path: Path, evidence_id: str, char_limit: int) -> dict[str, Any]:
        suffix = path.suffix.casefold()
        if suffix in _TEXT_SUFFIXES:
            content = path.read_text(encoding=self.text_encoding)
            clipped, truncated = _clip(content, char_limit)
            return {
                "evidence_id": evidence_id,
                "path": str(path),
                "kind": "text",
                "status": "loaded",
                "content": clipped,
                "truncated": truncated,
            }
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            parts: list[str] = []
            for page_number, page in enumerate(reader.pages, start=1):
                parts.append(f"[page {page_number}]\n{page.extract_text() or ''}")
                if sum(len(part) for part in parts) >= char_limit:
                    break
            clipped, truncated = _clip("\n\n".join(parts), char_limit)
            return {
                "evidence_id": evidence_id,
                "path": str(path),
                "kind": "document",
                "document_type": "pdf",
                "status": "loaded",
                "content": clipped,
                "truncated": truncated or len(parts) < len(reader.pages),
                "total_pages": len(reader.pages),
            }
        if suffix == ".docx":
            document = Document(str(path))
            content = "\n".join(paragraph.text for paragraph in document.paragraphs)
            clipped, truncated = _clip(content, char_limit)
            return {
                "evidence_id": evidence_id,
                "path": str(path),
                "kind": "document",
                "document_type": "docx",
                "status": "loaded",
                "content": clipped,
                "truncated": truncated,
                "total_paragraphs": len(document.paragraphs),
            }
        if suffix in _IMAGE_SUFFIXES:
            with Image.open(path) as image:
                image.verify()
            analysis = self.analyzer.analyze(
                path=path,
                prompt="Describe the image faithfully for use as evidence in the current user request.",
            )
            clipped, truncated = _clip(analysis, char_limit)
            return {
                "evidence_id": evidence_id,
                "path": str(path),
                "kind": "image",
                "status": "loaded",
                "model": self.analyzer.model,
                "content": clipped,
                "truncated": truncated,
            }
        return {
            "evidence_id": evidence_id,
            "path": str(path),
            "kind": "unsupported",
            "status": "unsupported_attachment_type",
            "suffix": suffix,
        }


def _clip(text: str, limit: int) -> tuple[str, bool]:
    value = str(text)
    if len(value) <= limit:
        return value, False
    return value[:limit], True
